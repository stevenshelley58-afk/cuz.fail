from __future__ import annotations

import csv
import json
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from docx import Document as DocxDocument
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.orm import Session

from draftcheck_compliance.service import LIABILITY_NOTICE, ComplianceService
from draftcheck_core.audit import record_audit
from draftcheck_core.json_utils import from_json, to_json
from draftcheck_core.models import Export, Project
from draftcheck_core.object_storage import get_export_storage
from draftcheck_document_ai.rfi import RfiService
from draftcheck_shared.schemas import ExportManifest

SOURCE_SIGNOFF_NOTICE = "Human signoff required before submission. This export is not submission-ready."


class ExportService:
    def __init__(self, db: Session):
        self.db = db
        self.storage = get_export_storage()

    def create_export(
        self,
        project_id: str,
        *,
        format: str,
        sections: list[str],
        created_by: str = "dev-user",
    ) -> ExportManifest:
        project = self.db.get(Project, project_id)
        if not project:
            raise KeyError("Project not found")
        payload = self._build_payload(project_id, sections)
        export_readiness = _export_readiness(payload)
        export = Export(
            project_id=project_id,
            export_type="response_pack",
            format=format,
            status="pending_human_signoff",
            created_by=created_by,
            manifest_json=to_json(
                {
                    "project_id": project_id,
                    "sections": sections,
                    "liability_notice": LIABILITY_NOTICE,
                    "requires_human_signoff": True,
                    "submission_ready": False,
                    "human_signoff_status": "required",
                    "source_signoff_notice": SOURCE_SIGNOFF_NOTICE,
                    "export_readiness": export_readiness,
                }
            ),
        )
        self.db.add(export)
        self.db.flush()

        filename = f"{export.id}.{_extension(format)}"
        content = self._render_export_bytes(format, payload, filename)
        stored = self.storage.put_bytes(f"exports/{filename}", content)
        export.object_key = stored.object_key
        export.file_sha256 = stored.content_sha256
        export.manifest_json = to_json(
            {
                **from_json(export.manifest_json, {}),
                "object_key": export.object_key,
                "file_sha256": export.file_sha256,
                "source_version_ids": _source_version_ids(payload),
            }
        )
        record_audit(
            self.db,
            action="export.created",
            target_type="export",
            target_id=export.id,
            actor_id=created_by,
            project_id=project_id,
            metadata={"format": format, "sections": sections, "file_sha256": export.file_sha256},
        )
        return _export_to_schema(export)

    def _render_export_bytes(self, format: str, payload: dict[str, Any], filename: str) -> bytes:
        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / filename
            if format == "json":
                path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
            elif format == "csv":
                self._write_csv(path, payload)
            elif format == "docx":
                self._write_docx(path, payload)
            elif format == "xlsx":
                self._write_xlsx(path, payload)
            elif format == "html":
                self._write_html(path, payload)
            else:
                raise ValueError("Unsupported export format")
            return path.read_bytes()

    def list_exports(self, project_id: str) -> list[ExportManifest]:
        rows = self.db.scalars(
            select(Export).where(Export.project_id == project_id).order_by(Export.created_at.desc())
        ).all()
        return [_export_to_schema(row) for row in rows]

    def get_export(self, project_id: str, export_id: str) -> ExportManifest:
        export = self.db.get(Export, export_id)
        if not export or export.project_id != project_id:
            raise KeyError("Export not found")
        return _export_to_schema(export)

    def export_file_path(self, project_id: str, export_id: str) -> Path:
        export = self.db.get(Export, export_id)
        if not export or export.project_id != project_id:
            raise KeyError("Export not found")
        if not export.object_key:
            raise FileNotFoundError("Export has no generated file")
        if export.object_key.startswith("s3://"):
            raise FileNotFoundError("Export file is stored remotely")
        path = Path(export.object_key)
        if not path.exists():
            raise FileNotFoundError("Export file is missing from object storage")
        return path

    def export_file_bytes(self, project_id: str, export_id: str) -> tuple[bytes, str]:
        export = self.db.get(Export, export_id)
        if not export or export.project_id != project_id:
            raise KeyError("Export not found")
        if not export.object_key:
            raise FileNotFoundError("Export has no generated file")
        if not self.storage.exists(export.object_key):
            raise FileNotFoundError("Export file is missing from object storage")
        return self.storage.get_bytes(export.object_key), Path(export.object_key).name

    def _build_payload(self, project_id: str, sections: list[str]) -> dict[str, Any]:
        compliance = ComplianceService(self.db)
        rfi = RfiService(self.db)
        payload: dict[str, Any] = {
            "project_id": project_id,
            "liability_notice": LIABILITY_NOTICE,
            "requires_human_signoff": True,
            "submission_ready": False,
            "human_signoff_status": "required",
            "source_signoff_notice": SOURCE_SIGNOFF_NOTICE,
            "sections": {},
        }
        if "compliance_matrix" in sections:
            payload["sections"]["compliance_matrix"] = [
                result.model_dump(mode="json") for result in compliance.list_latest_run_results(project_id)
            ]
        if "rfi_response" in sections:
            latest_response = rfi.latest_response(project_id)
            payload["sections"]["rfi_response"] = (
                [latest_response.model_dump(mode="json")] if latest_response else []
            )
        if "rfi_items" in sections:
            payload["sections"]["rfi_items"] = [
                item.model_dump(mode="json") for item in rfi.list_items(project_id)
            ]
        if "source_list" in sections:
            payload["sections"]["source_list"] = _source_version_ids(payload)
        if "assumptions" in sections:
            payload["sections"]["assumptions"] = _collect(payload, "assumptions")
        if "missing_info" in sections:
            payload["sections"]["missing_info"] = _collect(payload, "missing_information")
        payload["export_readiness"] = _export_readiness(payload)
        return payload

    def _write_csv(self, path: Path, payload: dict[str, Any]) -> None:
        rows = payload["sections"].get("compliance_matrix", [])
        with path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(
                handle,
                fieldnames=["check_key", "label", "category", "status", "proposed", "requirement"],
            )
            handle.write(f"{payload['liability_notice']}\n")
            handle.write(f"{payload['source_signoff_notice']}\n")
            writer.writeheader()
            for row in rows:
                writer.writerow({key: row.get(key, "") for key in writer.fieldnames or []})

    def _write_docx(self, path: Path, payload: dict[str, Any]) -> None:
        doc = DocxDocument()
        doc.add_heading("DraftCheck WA Response Pack", level=1)
        doc.add_paragraph(payload["liability_notice"])
        for response in payload["sections"].get("rfi_response", []):
            doc.add_heading(response["title"], level=2)
            doc.add_paragraph(response["draft_text"])
        doc.add_heading("Human Signoff", level=2)
        doc.add_paragraph("Required before submission.")
        doc.save(str(path))

    def _write_xlsx(self, path: Path, payload: dict[str, Any]) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "Compliance Matrix"
        ws.append(["Check", "Category", "Status", "Proposed", "Requirement"])
        for row in payload["sections"].get("compliance_matrix", []):
            ws.append(
                [
                    row.get("label"),
                    row.get("category"),
                    row.get("status"),
                    row.get("proposed"),
                    row.get("requirement"),
                ]
            )
        signoff = wb.create_sheet("Human Signoff")
        signoff.append(["Liability notice", payload["liability_notice"]])
        signoff.append(["Signoff notice", payload["source_signoff_notice"]])
        wb.save(str(path))

    def _write_html(self, path: Path, payload: dict[str, Any]) -> None:
        body = (
            f"<h1>DraftCheck WA Response Pack</h1><p>{LIABILITY_NOTICE}</p>"
            f"<h2>Human Signoff</h2><p>{payload['source_signoff_notice']}</p>"
        )
        path.write_text(f"<!doctype html><html><body>{body}</body></html>", encoding="utf-8")


def _extension(format: str) -> str:
    return {"json": "json", "csv": "csv", "docx": "docx", "xlsx": "xlsx", "html": "html"}[format]


def _source_version_ids(payload: dict[str, Any]) -> list[str]:
    ids: set[str] = set()
    for section in payload.get("sections", {}).values():
        if isinstance(section, list):
            for row in section:
                if not isinstance(row, dict):
                    if isinstance(row, str) and row.startswith("sv_"):
                        ids.add(row)
                    continue
                for citation in row.get("citations", []) or row.get("source_requirement_candidates", []):
                    source_version_id = citation.get("source_version_id")
                    if source_version_id:
                        ids.add(source_version_id)
    return sorted(ids)


def _collect(payload: dict[str, Any], key: str) -> list[str]:
    values: set[str] = set()
    for section in payload.get("sections", {}).values():
        if isinstance(section, list):
            for row in section:
                for value in row.get(key, []) or []:
                    values.add(value)
    return sorted(values)


_BLOCKING_COMPLIANCE_STATUSES = {"missing_info", "needs_human_review", "unsupported", "stale"}


def _export_readiness(payload: dict[str, Any]) -> dict[str, Any]:
    blocking_issues: list[dict[str, Any]] = []
    compliance_rows = payload.get("sections", {}).get("compliance_matrix", [])
    if isinstance(compliance_rows, list):
        for row in compliance_rows:
            if not isinstance(row, dict):
                continue
            status = row.get("status")
            if status in _BLOCKING_COMPLIANCE_STATUSES:
                blocking_issues.append(
                    {
                        "section": "compliance_matrix",
                        "target_id": row.get("id"),
                        "check_key": row.get("check_key"),
                        "status": status,
                        "reason": "Exported compliance result is not a resolved likely pass/fail/not-applicable result.",
                    }
                )

    drafts = payload.get("sections", {}).get("rfi_response", [])
    if isinstance(drafts, list):
        for draft in drafts:
            if not isinstance(draft, dict):
                continue
            missing_information = draft.get("missing_information") or []
            if draft.get("requires_human_review") or missing_information:
                blocking_issues.append(
                    {
                        "section": "rfi_response",
                        "target_id": draft.get("id"),
                        "status": "needs_human_review",
                        "reason": "Exported response draft still requires human review or missing information.",
                    }
                )

    rfi_items = payload.get("sections", {}).get("rfi_items", [])
    if isinstance(rfi_items, list):
        for item in rfi_items:
            if not isinstance(item, dict):
                continue
            missing_evidence = item.get("missing_evidence") or []
            status = item.get("status")
            if missing_evidence or status in {"open", "needs_review", "needs_human_review", "unsupported"}:
                blocking_issues.append(
                    {
                        "section": "rfi_items",
                        "target_id": item.get("id"),
                        "status": status or "needs_human_review",
                        "reason": "Exported RFI item still has missing evidence or unresolved review status.",
                    }
                )

    return {
        "status": "blocked" if blocking_issues else "ready",
        "blocking_issue_count": len(blocking_issues),
        "blocking_issues": blocking_issues,
    }


def _export_to_schema(row: Export) -> ExportManifest:
    return ExportManifest(
        id=row.id,
        project_id=row.project_id,
        export_type=row.export_type,
        format=row.format,
        status=row.status,
        object_key=row.object_key,
        file_sha256=row.file_sha256,
        manifest=from_json(row.manifest_json, {}),
        created_at=row.created_at,
    )
