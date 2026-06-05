from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from math import hypot

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(frozen=True)
class DxfEntity:
    type: str
    fields: dict[str, list[str]]


def extract_text_from_bytes(content: bytes, content_type: str) -> str:
    return "\n".join(extract_pages_from_bytes(content, content_type)).strip()


def extract_pages_from_bytes(content: bytes, content_type: str) -> list[str]:
    content_hint = content_type.lower()
    if "dxf" in content_hint or ("filename=" in content_hint and content_hint.endswith(".dxf")):
        return [_extract_dxf_summary(content)]
    if "pdf" in content_type:
        reader = PdfReader(BytesIO(content))
        return [(page.extract_text() or "").strip() for page in reader.pages] or [""]
    if "wordprocessingml" in content_type or "msword" in content_type:
        document = DocxDocument(BytesIO(content))
        paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in document.tables
            for row in table.rows
        ]
        text = "\n".join(part for part in [*paragraph_text, *table_text] if part).strip()
        return [text]
    if "html" in content_type:
        soup = BeautifulSoup(content.decode("utf-8", errors="ignore"), "html.parser")
        return [soup.get_text("\n", strip=True)]
    text = content.decode("utf-8", errors="ignore")
    pages = [page.strip() for page in text.split("\f")]
    return pages if any(pages) else [text.strip()]


def _extract_dxf_summary(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    lines = [line.strip() for line in text.splitlines()]
    pairs = list(zip(lines[0::2], lines[1::2], strict=False))
    units = _dxf_units(pairs)
    entities = _dxf_entities(pairs)
    entity_counts = Counter(entity.type for entity in entities)
    layers: set[str] = set()
    drawing_text: list[str] = []
    coordinates: list[float] = []
    dimension_summaries: list[str] = []

    for entity in entities:
        fields = entity.fields
        layer = _first(fields, "8")
        if layer:
            layers.add(layer)
        for value in [*_values(fields, "1"), *_values(fields, "3")]:
            cleaned = _clean_dxf_text(value)
            if cleaned and any(char.isalnum() for char in cleaned):
                drawing_text.append(cleaned)
        coordinates.extend(_entity_coordinates(fields))
        dimension_summaries.extend(_entity_dimension_summaries(entity, units))

    bbox = ""
    if coordinates:
        bbox = f"Coordinate range {min(coordinates):g} to {max(coordinates):g}."
    entity_summary = ", ".join(f"{name}:{entity_counts[name]}" for name in sorted(entity_counts))
    layer_summary = ", ".join(sorted(layers)[:50])
    text_summary = "\n".join([*dimension_summaries[:200], *drawing_text[:200]])
    return "\n".join(
        part
        for part in [
            "DXF drawing extraction summary.",
            f"DXF declared units: {units}." if units else "",
            f"Entities: {entity_summary}." if entity_summary else "",
            f"Layers: {layer_summary}." if layer_summary else "",
            bbox,
            text_summary,
        ]
        if part
    )


def _dxf_units(pairs: list[tuple[str, str]]) -> str | None:
    units_by_code = {
        "0": "unitless",
        "1": "in",
        "2": "ft",
        "4": "mm",
        "5": "cm",
        "6": "m",
    }
    for index, (code, value) in enumerate(pairs):
        if code == "9" and value.upper() == "$INSUNITS":
            for next_code, next_value in pairs[index + 1 : index + 5]:
                if next_code == "70":
                    return units_by_code.get(next_value)
    return None


def _dxf_entities(pairs: list[tuple[str, str]]) -> list[DxfEntity]:
    entity_names = {
        "LINE",
        "LWPOLYLINE",
        "POLYLINE",
        "CIRCLE",
        "ARC",
        "TEXT",
        "MTEXT",
        "INSERT",
        "DIMENSION",
    }
    entities: list[DxfEntity] = []
    current: DxfEntity | None = None
    for code, value in pairs:
        if code == "0":
            if current:
                entities.append(current)
            current = DxfEntity(value, {}) if value in entity_names else None
            continue
        if current is not None:
            current.fields.setdefault(code, []).append(value)
    if current:
        entities.append(current)
    return entities


def _entity_dimension_summaries(entity: DxfEntity, units: str | None) -> list[str]:
    entity_type = entity.type
    fields = entity.fields
    layer = _first(fields, "8") or "unknown"
    if entity_type == "LINE":
        start_x = _first_float(fields, "10")
        start_y = _first_float(fields, "20")
        end_x = _first_float(fields, "11")
        end_y = _first_float(fields, "21")
        if start_x is not None and start_y is not None and end_x is not None and end_y is not None:
            length = hypot(end_x - start_x, end_y - start_y)
            return [f"Line length on layer {layer}: {_format_dxf_measure(length, units)}."]
    if entity_type == "LWPOLYLINE":
        xs = _float_values(fields, "10")
        ys = _float_values(fields, "20")
        if len(xs) >= 2 and len(xs) == len(ys):
            length = sum(hypot(xs[index] - xs[index - 1], ys[index] - ys[index - 1]) for index in range(1, len(xs)))
            closed_flag = int(_first_float(fields, "70") or 0)
            if closed_flag & 1:
                length += hypot(xs[0] - xs[-1], ys[0] - ys[-1])
            return [f"Polyline length on layer {layer}: {_format_dxf_measure(length, units)}."]
    if entity_type == "DIMENSION":
        actual = _first_float(fields, "42")
        label = _clean_dxf_text(_first(fields, "1") or "")
        if actual is not None:
            text = f"Dimension measurement on layer {layer}: {_format_dxf_measure(actual, units)}."
            return [f"{text} Label: {label}." if label else text]
    return []


def _entity_coordinates(fields: dict[str, list[str]]) -> list[float]:
    values: list[float] = []
    for code in {"10", "11", "20", "21", "30", "31", "40", "41", "42"}:
        values.extend(_float_values(fields, code))
    return values


def _format_dxf_measure(value: float, units: str | None) -> str:
    if units in {"mm", "m", "cm", "in", "ft"}:
        return f"{value:g}{units}"
    return f"{value:g} drawing units"


def _first(fields: dict[str, list[str]], code: str) -> str | None:
    values = fields.get(code) or []
    return values[0] if values else None


def _values(fields: dict[str, list[str]], code: str) -> list[str]:
    return fields.get(code) or []


def _first_float(fields: dict[str, list[str]], code: str) -> float | None:
    values = _float_values(fields, code)
    return values[0] if values else None


def _float_values(fields: dict[str, list[str]], code: str) -> list[float]:
    values: list[float] = []
    for value in fields.get(code) or []:
        try:
            values.append(float(value))
        except ValueError:
            pass
    return values


def _clean_dxf_text(value: str) -> str:
    return value.replace("\\P", "\n").replace("{", "").replace("}", "").strip()
