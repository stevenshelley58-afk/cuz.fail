"""V3 document upload and parser router.

Endpoints
---------
  GET  /documents/parsers                          → parser capability list
  GET  /documents/parsers/accuracy                 → accuracy report

  POST /documents/upload                           → upload + fact extraction
  GET  /documents/{doc_id}/facts                   → list extracted facts
  PATCH /documents/{doc_id}/facts/{fact_id}        → update value or status
  POST  /documents/{doc_id}/facts/{fact_id}/promote → promote to PropertyFact

The upload endpoint stores files content-addressed at
  /srv/draftcheck/storage/{sha256[:2]}/{sha256}
and creates Document + DocumentPage + DocumentFact rows in the database.

When DATABASE_URL is absent the in-memory library is used for the parser
capability and accuracy report endpoints.  The upload / facts endpoints
return 503 when the database is not configured.
"""

from __future__ import annotations

import hashlib
import io
import os
import uuid
from collections.abc import Generator
from pathlib import Path, PurePath
from typing import Annotated, Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.encoders import jsonable_encoder
from pydantic import BaseModel, ConfigDict, Field
from sqlalchemy.orm import Session

from draftcheck.api.auth import get_current_session, require_allowed_origin
from draftcheck.db.engine import create_session_factory
from draftcheck.db.models import (
    Document,
    DocumentFact as OrmDocumentFact,
    DocumentPage as OrmDocumentPage,
    PropertyFact,
)
from draftcheck.domain.documents import (
    DocumentFact,
    DocumentNotFoundError,
    DocumentParser,
    DocumentReviewStatus,
    InMemoryDocumentLibrary,
    sample_parser_accuracy_report,
)
from draftcheck.domain.documents.facts import DocumentFactService
from draftcheck.domain.identity import ActiveSession

router = APIRouter(tags=["documents"])

_document_library: InMemoryDocumentLibrary | None = None
_fact_service = DocumentFactService()

STORAGE_ROOT = Path(os.getenv("DRAFTCHECK_STORAGE_ROOT", "/srv/draftcheck/storage"))
MAX_UPLOAD_BYTES = 15 * 1024 * 1024  # 15 MB


# ---------------------------------------------------------------------------
# In-memory library (parser capability / accuracy report only)
# ---------------------------------------------------------------------------


def get_document_library() -> InMemoryDocumentLibrary:
    global _document_library
    if _document_library is None:
        _document_library = InMemoryDocumentLibrary()
    return _document_library


# ---------------------------------------------------------------------------
# DB session dependency (mirrors projects.py)
# ---------------------------------------------------------------------------


def get_db_session() -> Generator[Session, None, None]:
    database_url = os.getenv("DATABASE_URL")
    if not database_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="DATABASE_URL is not configured; durable document storage unavailable.",
        )
    factory = create_session_factory(database_url)
    db: Session = factory()
    try:
        yield db
        db.commit()
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


DbSession = Annotated[Session, Depends(get_db_session)]


# ---------------------------------------------------------------------------
# Request / response schemas
# ---------------------------------------------------------------------------


class FactUpdateRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    numeric_value: float | None = Field(default=None)
    status: str | None = Field(default=None, pattern=r"^(pending_review|confirmed|rejected)$")


class ReviewDocumentFactPayload(BaseModel):
    model_config = ConfigDict(extra="forbid")

    review_status: DocumentReviewStatus
    note: str | None = Field(default=None, max_length=1000)


# ---------------------------------------------------------------------------
# Serialisation helpers
# ---------------------------------------------------------------------------


def _orm_fact_payload(fact: OrmDocumentFact) -> dict[str, Any]:
    v = fact.value_json or {}
    return {
        "fact_id": str(fact.id),
        "fact_key": fact.check_key,
        "fact_kind": fact.fact_kind,
        "numeric_value": v.get("numeric_value"),
        "unit": v.get("unit"),
        "source_text": v.get("source_text"),
        "page_number": v.get("page_number"),
        "confidence": fact.confidence,
        "review_status": fact.review_status,
        "promoted_to_measurement": fact.promoted_to_measurement,
    }


def _inmem_fact_payload(fact: DocumentFact) -> dict[str, Any]:
    payload = jsonable_encoder(fact)
    payload["review_status"] = fact.review_status.value
    return payload


def _not_found(exc: DocumentNotFoundError) -> HTTPException:
    return HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Document item not found: {exc}")


# ---------------------------------------------------------------------------
# Storage helpers
# ---------------------------------------------------------------------------


def _store_content(content: bytes, sha256: str) -> Path:
    """Write bytes to content-addressed storage; return the file path."""
    dest = STORAGE_ROOT / sha256[:2] / sha256
    dest.parent.mkdir(parents=True, exist_ok=True)
    if not dest.exists():
        dest.write_bytes(content)
    return dest


def _extract_text_from_content(media_type: str, filename: str, content: bytes) -> tuple[list[str], str]:
    """Return (per-page texts, parser_name).

    Uses pypdf for PDF, python-docx for DOCX, plain text for everything else.
    DXF is treated as text-only metadata.
    """
    suffix = PurePath(filename).suffix.lower()
    parser_name = "draftcheck.plain_text_parser"

    if media_type == "application/pdf" or suffix == ".pdf":
        parser_name = "draftcheck.pdf_text_parser"
        pages = _pdf_pages(content)
        return pages, parser_name

    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or suffix == ".docx":
        parser_name = "draftcheck.docx_text_parser"
        text = _docx_text(content)
        return [text] if text.strip() else [], parser_name

    if suffix == ".dxf" or "dxf" in media_type:
        parser_name = "draftcheck.dxf_text_parser"
        # DXF: decode as text for metadata / entity labels only
        text = _decode_bytes(content)
        return [text] if text.strip() else [], parser_name

    # TXT / CSV / anything else
    text = _decode_bytes(content)
    return [text] if text.strip() else [], parser_name


def _pdf_pages(content: bytes) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return [_decode_bytes(content)]
    try:
        reader = PdfReader(io.BytesIO(content))
        return [page.extract_text() or "" for page in reader.pages]
    except Exception:
        return [_decode_bytes(content)]


def _docx_text(content: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return _decode_bytes(content)
    try:
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return _decode_bytes(content)


def _decode_bytes(content: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _safe_filename(filename: str) -> str:
    import re

    name = PurePath(filename or "upload.bin").name.replace("\x00", "")
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" .")
    return name or "upload.bin"


def _infer_media_type(upload: UploadFile) -> str:
    if upload.content_type:
        return upload.content_type.split(";")[0].strip().lower()
    suffix = PurePath(upload.filename or "").suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".dxf": "application/dxf",
        ".csv": "text/csv",
    }.get(suffix, "application/octet-stream")


# ---------------------------------------------------------------------------
# Parser capability / accuracy endpoints (in-memory, no DB needed)
# ---------------------------------------------------------------------------


@router.get("/documents/parsers", tags=["documents"])
def list_document_parsers() -> dict[str, Any]:
    capabilities = DocumentParser().capabilities()
    return {
        "items": jsonable_encoder(capabilities),
        "count": len(capabilities),
        "accuracy_gate": {
            "status": "not_beta_ready",
            "reason": "parser coverage and golden eval accuracy gates have not passed",
            "required_before_beta": [
                "real PDF/DOCX/DXF/IFC fixture set",
                "per-field precision/recall report",
                "automated review gate",
                "no raster measurement without calibration",
            ],
        },
    }


@router.get("/documents/parsers/accuracy", tags=["documents"])
def get_document_parser_accuracy() -> dict[str, Any]:
    return sample_parser_accuracy_report()


# ---------------------------------------------------------------------------
# POST /documents/upload
# ---------------------------------------------------------------------------


@router.post("/documents/upload", tags=["documents"])
async def upload_document(
    file: Annotated[UploadFile, File()],
    project_id: str,
    db: DbSession,
    _allowed_origin: Annotated[None, Depends(require_allowed_origin)],
    active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    """Accept a multipart file upload, store it, parse text, extract facts.

    Returns document_id and extracted_facts[].
    """
    if active_session.org is None:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Authenticated org required.")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="File is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="File exceeds the 15 MB upload limit.",
        )

    filename = _safe_filename(file.filename or "upload.bin")
    media_type = _infer_media_type(file)
    sha256 = hashlib.sha256(content).hexdigest()
    storage_path = str(_store_content(content, sha256))

    # --- Create Document row ---
    doc_id = uuid.uuid4()
    org_id = active_session.org.id

    try:
        proj_uuid = uuid.UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid project_id UUID.") from exc

    document = Document(
        id=doc_id,
        org_id=org_id,
        project_id=proj_uuid,
        uploaded_by_user_id=active_session.user.id if active_session.user else None,
        title=filename,
        document_type=_doc_type_from_media(media_type, filename),
        status="uploaded",
        storage_path=storage_path,
        sha256=sha256,
        media_type=media_type,
        size_bytes=len(content),
        metadata_json={"original_filename": filename},
    )
    db.add(document)
    db.flush()  # get the id before adding pages/facts

    # --- Extract text + create DocumentPage rows ---
    page_texts, parser_name = _extract_text_from_content(media_type, filename, content)
    orm_pages: list[OrmDocumentPage] = []
    for page_number, page_text in enumerate(page_texts, start=1):
        page = OrmDocumentPage(
            id=uuid.uuid4(),
            document_id=doc_id,
            page_number=page_number,
            text=page_text,
            metadata_json={"parser_name": parser_name, "parser_version": DocumentFactService.PARSER_VERSION},
        )
        db.add(page)
        orm_pages.append(page)

    db.flush()

    # --- Extract facts ---
    all_orm_facts: list[OrmDocumentFact] = []
    for page_row in orm_pages:
        page_facts = _fact_service.extract_facts_from_text(
            text=page_row.text or "",
            document_id=doc_id,
            page_number=page_row.page_number,
            org_id=org_id,
            project_id=proj_uuid,
            page_id=page_row.id,
        )
        for fact in page_facts:
            db.add(fact)
            all_orm_facts.append(fact)

    db.flush()

    return {
        "document_id": str(doc_id),
        "filename": filename,
        "media_type": media_type,
        "size_bytes": len(content),
        "sha256": sha256,
        "page_count": len(orm_pages),
        "extracted_facts": [_orm_fact_payload(f) for f in all_orm_facts],
        "fact_count": len(all_orm_facts),
        "review_required": True,
        "advisory_notice": "All extracted measurements are advisory. Promote facts to confirm before running compliance.",
    }


def _doc_type_from_media(media_type: str, filename: str) -> str:
    suffix = PurePath(filename).suffix.lower()
    mapping = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/dxf": "dxf",
        "text/plain": "txt",
        "text/csv": "csv",
    }
    result = mapping.get(media_type)
    if result:
        return result
    return suffix.lstrip(".") or "unknown"


# ---------------------------------------------------------------------------
# PATCH /documents/{doc_id}/facts/{fact_id}
# ---------------------------------------------------------------------------


@router.patch("/documents/{doc_id}/facts/{fact_id}", tags=["documents"])
def update_document_fact(
    doc_id: str,
    fact_id: str,
    payload: FactUpdateRequest,
    db: DbSession,
    _allowed_origin: Annotated[None, Depends(require_allowed_origin)],
    _active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    """Update a fact's numeric_value and/or status."""
    fact = _get_fact_or_404(db, doc_id, fact_id)

    if payload.numeric_value is not None:
        updated_value = dict(fact.value_json or {})
        updated_value["numeric_value"] = payload.numeric_value
        fact.value_json = updated_value

    if payload.status is not None:
        fact.review_status = payload.status

    db.flush()
    return _orm_fact_payload(fact)


# ---------------------------------------------------------------------------
# POST /documents/{doc_id}/facts/{fact_id}/promote
# ---------------------------------------------------------------------------


@router.post("/documents/{doc_id}/facts/{fact_id}/promote", tags=["documents"])
def promote_document_fact(
    doc_id: str,
    fact_id: str,
    db: DbSession,
    _allowed_origin: Annotated[None, Depends(require_allowed_origin)],
    active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    """Set status='confirmed', promote fact to PropertyFact on the project.

    Creates a new PropertyFact row linked to the document fact's project.
    """
    fact = _get_fact_or_404(db, doc_id, fact_id)
    fact.review_status = "confirmed"
    fact.promoted_to_measurement = True

    v = fact.value_json or {}
    numeric_value = v.get("numeric_value")
    unit = v.get("unit", "")
    fact_key = fact.check_key or fact.fact_kind

    property_fact = PropertyFact(
        id=uuid.uuid4(),
        org_id=fact.org_id,
        project_id=fact.project_id,
        fact_type=fact_key,
        value_json={"value": numeric_value, "unit": unit, "document_fact_id": str(fact.id)},
        confidence=fact.confidence,
        method="document_extraction_promoted",
        provenance_json={
            "entered_by": str(active_session.user.id) if active_session.user else "system",
            "reason": "promoted from document fact",
            "method": "document_extraction_promoted",
            "source_document_id": str(fact.document_id),
            "source_fact_id": str(fact.id),
        },
        review_status="confirmed",
    )
    db.add(property_fact)
    db.flush()

    return {
        "fact_id": str(fact.id),
        "review_status": fact.review_status,
        "promoted_to_measurement": fact.promoted_to_measurement,
        "property_fact_id": str(property_fact.id),
        "advisory_notice": "Promoted fact is advisory. Not a legal or compliance certification.",
    }


# ---------------------------------------------------------------------------
# Legacy in-memory endpoints (kept for backward compatibility with existing tests)
# ---------------------------------------------------------------------------


@router.post("/documents/projects/{project_id}/upload", tags=["documents"])
async def upload_project_document(
    project_id: str,
    file: Annotated[UploadFile, File()],
    library: Annotated[InMemoryDocumentLibrary, Depends(get_document_library)],
    _allowed_origin: Annotated[None, Depends(require_allowed_origin)],
    active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    content = await file.read()
    try:
        result = library.upload(
            org_id=str(active_session.org.id),
            project_id=project_id,
            user_id=str(active_session.user.id),
            filename=file.filename or "upload.bin",
            media_type=file.content_type or "",
            content=content,
        )
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail=str(exc)) from exc
    return {
        "document": jsonable_encoder(result.document),
        "pages": len(result.pages),
        "facts": [_inmem_fact_payload(fact) for fact in result.facts],
        "fact_count": len(result.facts),
        "review_required": True,
    }


@router.get("/documents/{document_id}", tags=["documents"])
def get_document(
    document_id: str,
    library: Annotated[InMemoryDocumentLibrary, Depends(get_document_library)],
    _active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    try:
        return jsonable_encoder(library.get_document(document_id))
    except DocumentNotFoundError as exc:
        raise _not_found(exc) from exc


@router.get("/documents/{document_id}/pages", tags=["documents"])
def get_document_pages(
    document_id: str,
    library: Annotated[InMemoryDocumentLibrary, Depends(get_document_library)],
    _active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    try:
        pages = library.get_pages(document_id)
    except DocumentNotFoundError as exc:
        raise _not_found(exc) from exc
    return {"items": jsonable_encoder(pages), "count": len(pages)}


@router.get("/documents/{document_id}/facts", tags=["documents"])
def get_document_facts_inmem(
    document_id: str,
    library: Annotated[InMemoryDocumentLibrary, Depends(get_document_library)],
    _active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    try:
        facts = library.get_facts(document_id)
    except DocumentNotFoundError as exc:
        raise _not_found(exc) from exc
    return {"items": [_inmem_fact_payload(f) for f in facts], "count": len(facts)}


@router.post("/documents/{document_id}/facts/{fact_id}/review", tags=["documents"])
def review_document_fact(
    document_id: str,
    fact_id: str,
    payload: ReviewDocumentFactPayload,
    library: Annotated[InMemoryDocumentLibrary, Depends(get_document_library)],
    _allowed_origin: Annotated[None, Depends(require_allowed_origin)],
    active_session: Annotated[ActiveSession, Depends(get_current_session)],
) -> dict[str, Any]:
    try:
        fact = library.review_fact(
            document_id=document_id,
            fact_id=fact_id,
            review_status=payload.review_status,
            reviewed_by=str(active_session.user.id),
            note=payload.note,
        )
    except DocumentNotFoundError as exc:
        raise _not_found(exc) from exc
    return _inmem_fact_payload(fact)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _get_fact_or_404(db: Session, doc_id: str, fact_id: str) -> OrmDocumentFact:
    try:
        doc_uuid = uuid.UUID(doc_id)
        fact_uuid = uuid.UUID(fact_id)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid UUID.") from exc

    fact: OrmDocumentFact | None = (
        db.query(OrmDocumentFact)
        .filter(
            OrmDocumentFact.id == fact_uuid,
            OrmDocumentFact.document_id == doc_uuid,
        )
        .first()
    )
    if fact is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Fact {fact_id} not found on document {doc_id}.",
        )
    return fact
