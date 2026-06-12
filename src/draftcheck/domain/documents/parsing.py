"""Conservative in-memory document parser for V3 document contract endpoints."""

from __future__ import annotations

import csv
import hashlib
import io
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from enum import StrEnum
from pathlib import PurePath
from typing import Any
from uuid import uuid4

from draftcheck.domain.documents.chunks import (
    DocumentChunk,
    DocumentChunkSearchHit,
    build_document_chunks,
    search_document_chunks,
)
from draftcheck.domain.documents.parsers import (
    DocumentArtifact,
    DocumentContentParser,
    DocumentParserRegistry,
    ParsedDocument,
)
from draftcheck.domain.sources.models import EmbeddingConfig


MAX_DOCUMENT_BYTES = 15 * 1024 * 1024
TEXT_MEDIA_TYPES = {
    "text/plain",
    "text/csv",
    "text/markdown",
    "application/json",
    "application/xml",
    "text/xml",
}
DXF_MEDIA_TYPES = {"application/dxf", "application/x-dxf", "image/vnd.dxf"}
DOCX_MEDIA_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
PDF_MEDIA_TYPES = {"application/pdf"}
IMAGE_PREFIXES = ("image/",)
IFC_MEDIA_TYPES = {"application/ifc", "application/x-step", "model/ifc"}
SAMPLE_PACK_CONTENT = b"""Lot area: 450 m2
Footprint: 218 m2
Open space: 180 m2
Front setback: 4.5 m
Garage width: 5.4 m
Boundary wall length: 0 m
"""
SAMPLE_EXPECTED_FACTS = {
    "lot area": (450.0, "m2"),
    "building footprint": (218.0, "m2"),
    "open space": (180.0, "m2"),
    "front setback": (4.5, "m"),
    "garage width": (5.4, "m"),
    "boundary wall length": (0.0, "m"),
}


class DocumentReviewStatus(StrEnum):
    PENDING_REVIEW = "pending_review"
    HUMAN_CONFIRMED = "human_confirmed"
    REJECTED = "rejected"


@dataclass(frozen=True)
class ParserCapability:
    media_type: str
    parser_name: str
    support_status: str
    notes: str


@dataclass(frozen=True)
class DocumentPage:
    id: str
    document_id: str
    page_number: int
    text: str
    parser_name: str
    parser_version: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class PdfTextBlock:
    text: str
    bbox: tuple[float, float, float, float]
    block_number: int | None


@dataclass(frozen=True)
class PdfVectorPath:
    bbox: tuple[float, float, float, float] | None
    item_count: int
    drawing_type: str | None


@dataclass(frozen=True)
class PdfPageExtraction:
    page_number: int
    text: str
    width: float | None
    height: float | None
    rotation_degrees: float | None
    text_blocks: tuple[PdfTextBlock, ...]
    vector_paths: tuple[PdfVectorPath, ...]
    extraction_method: str


@dataclass(frozen=True)
class DocumentFact:
    id: str
    document_id: str
    fact_type: str
    label: str
    value: Any
    numeric_value: float | None
    unit: str | None
    confidence: float
    review_status: DocumentReviewStatus
    source: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class DocumentRecord:
    id: str
    org_id: str
    project_id: str
    filename: str
    media_type: str
    size_bytes: int
    sha256: str
    parser_name: str
    parser_version: str
    parse_status: str
    uploaded_by_user_id: str
    uploaded_at: datetime
    metadata: dict[str, Any]


@dataclass(frozen=True)
class DocumentParseResult:
    document: DocumentRecord
    pages: tuple[DocumentPage, ...]
    chunks: tuple[DocumentChunk, ...]
    facts: tuple[DocumentFact, ...]


@dataclass(frozen=True)
class _DxfDimensionCandidate:
    measurement: float
    handle: str | None
    layer: str | None
    entity_type: str
    text_override: str | None = None
    text_override_numeric_value: float | None = None
    text_override_differs: bool = False
    block_name: str | None = None
    cad_space: str | None = None
    layout_name: str | None = None
    insert_handle: str | None = None
    insert_layer: str | None = None
    insert_scale_x: float | None = None
    insert_scale_y: float | None = None
    insert_scale_z: float | None = None
    insert_scale_applied: bool = False
    insert_scale_uncertain: bool = False
    scale_review_reason: str | None = None
    extraction_method: str = "dxf_group_code_dimension_entity"


@dataclass(frozen=True)
class _DxfUnitContext:
    code: int | None
    declared_units: str
    output_unit: str
    scale_to_metres: float | None
    confidence_cap: float | None

    @property
    def unit_declared(self) -> bool:
        return self.code is not None and self.code != 0 and self.scale_to_metres is not None


class DocumentNotFoundError(KeyError):
    """Raised when a document is unknown to the in-memory library."""


class DocumentParseError(ValueError):
    """Raised when content cannot be parsed by its declared format."""


class _RegisteredContentParser:
    def __init__(
        self,
        *,
        name: str,
        version: str,
        media_types: set[str] | None = None,
        suffixes: set[str] | None = None,
        media_prefixes: tuple[str, ...] = (),
    ) -> None:
        self.name = name
        self.version = version
        self.media_types = media_types or set()
        self.suffixes = suffixes or set()
        self.media_prefixes = media_prefixes

    def can_parse(self, media_type: str, filename: str) -> bool:
        suffix = PurePath(filename).suffix.lower()
        return (
            media_type in self.media_types
            or suffix in self.suffixes
            or any(media_type.startswith(prefix) for prefix in self.media_prefixes)
        )

    def parse(
        self,
        *,
        document_id: str,
        filename: str,
        media_type: str,
        content: bytes,
    ) -> ParsedDocument:
        text, extraction_notes = _extract_text(media_type, filename, content)
        pages = _pages(document_id, text, self.name, self.version, extraction_notes)
        facts = _facts(document_id, text, self.name, media_type)
        artifacts = _artifacts(
            parser_name=self.name,
            media_type=media_type,
            extraction_notes=extraction_notes,
            pages=pages,
            facts=facts,
        )
        parse_status = "parsed" if text.strip() or facts else "needs_more_info"
        metadata = {
            "extraction_notes": extraction_notes,
            "measurement_policy": "Measurements are advisory pending automated validation.",
            "raster_measurement_policy": "Raster/PDF/image measurements are not compliance-ready without calibration.",
            "artifacts": [artifact.to_dict() for artifact in artifacts],
        }
        return ParsedDocument(
            media_type=media_type,
            parser_name=self.name,
            parser_version=self.version,
            parse_status=parse_status,
            pages=pages,
            facts=facts,
            artifacts=artifacts,
            metadata=metadata,
        )


class _FallbackTextParser(_RegisteredContentParser):
    def can_parse(self, media_type: str, filename: str) -> bool:
        return True


class _PdfContentParser(_RegisteredContentParser):
    def parse(
        self,
        *,
        document_id: str,
        filename: str,
        media_type: str,
        content: bytes,
    ) -> ParsedDocument:
        extraction_notes = ["pdf_text_only_no_raster_measurements"]
        layouts = extract_pdf_page_layouts(content)
        pages = tuple(
            DocumentPage(
                id=f"page_{document_id}_{page.page_number}",
                document_id=document_id,
                page_number=page.page_number,
                text=page.text,
                parser_name=self.name,
                parser_version=self.version,
                metadata={
                    "notes": extraction_notes,
                    "width": page.width,
                    "height": page.height,
                    "rotation_degrees": page.rotation_degrees,
                    "extraction_method": page.extraction_method,
                    "text_blocks": [
                        {
                            "text": block.text,
                            "bbox": list(block.bbox),
                            "block_number": block.block_number,
                            "measurement_compliance_ready": False,
                            "measurement_readiness_reason": "pdf text block bbox is not a calibrated measurement",
                        }
                        for block in page.text_blocks
                    ],
                    "vector_paths": [
                        {
                            "bbox": list(path.bbox) if path.bbox else None,
                            "item_count": path.item_count,
                            "drawing_type": path.drawing_type,
                            "measurement_compliance_ready": False,
                            "measurement_readiness_reason": "pdf vector path is not a calibrated measurement",
                            "calibration_required": True,
                        }
                        for path in page.vector_paths
                    ],
                    "raster_measurement_policy": "PDF measurements require explicit calibration before promotion.",
                },
            )
            for page in layouts
        )
        text = "\n\n".join(page.text for page in layouts)
        facts = _facts(document_id, text, self.name, media_type)
        artifacts = _artifacts(
            parser_name=self.name,
            media_type=media_type,
            extraction_notes=extraction_notes,
            pages=pages,
            facts=facts,
        )
        parse_status = "parsed" if any(page.text.strip() for page in pages) or facts else "needs_more_info"
        metadata = {
            "extraction_notes": extraction_notes,
            "measurement_policy": "Measurements are advisory pending automated validation.",
            "raster_measurement_policy": "Raster/PDF/image measurements are not compliance-ready without calibration.",
            "artifacts": [artifact.to_dict() for artifact in artifacts],
        }
        return ParsedDocument(
            media_type=media_type,
            parser_name=self.name,
            parser_version=self.version,
            parse_status=parse_status,
            pages=pages,
            facts=facts,
            artifacts=artifacts,
            metadata=metadata,
        )


def _default_parser_registry(version: str) -> DocumentParserRegistry:
    parsers: tuple[DocumentContentParser, ...] = (
        _RegisteredContentParser(
            name="draftcheck.dxf_text_parser",
            version=version,
            media_types=DXF_MEDIA_TYPES,
            suffixes={".dxf"},
        ),
        _PdfContentParser(
            name="draftcheck.pdf_text_parser",
            version=version,
            media_types=PDF_MEDIA_TYPES,
            suffixes={".pdf"},
        ),
        _RegisteredContentParser(
            name="draftcheck.docx_text_parser",
            version=version,
            media_types=DOCX_MEDIA_TYPES,
            suffixes={".docx"},
        ),
        _RegisteredContentParser(
            name="draftcheck.ifc_text_parser",
            version=version,
            media_types=IFC_MEDIA_TYPES,
            suffixes={".ifc"},
        ),
        _RegisteredContentParser(
            name="draftcheck.image_ocr_queue",
            version=version,
            media_prefixes=IMAGE_PREFIXES,
        ),
        _RegisteredContentParser(
            name="draftcheck.csv_text_parser",
            version=version,
            media_types={"text/csv"},
            suffixes={".csv"},
        ),
        _FallbackTextParser(
            name="draftcheck.plain_text_parser",
            version=version,
            media_types=TEXT_MEDIA_TYPES,
            suffixes={".txt", ".md", ".json", ".xml"},
        ),
    )
    return DocumentParserRegistry(parsers)


class DocumentParser:
    """Extract reviewable document facts without producing compliance verdicts."""

    version = "v0.1"

    def __init__(self, registry: DocumentParserRegistry | None = None) -> None:
        self.registry = registry or _default_parser_registry(self.version)

    def capabilities(self) -> tuple[ParserCapability, ...]:
        return (
            ParserCapability("text/plain", "plain_text_parser", "active", "Text and simple numeric facts."),
            ParserCapability("text/csv", "csv_text_parser", "active", "CSV rows are flattened to text."),
            ParserCapability("application/dxf", "dxf_text_parser", "active", "DXF-like text dimensions and layers."),
            ParserCapability("application/pdf", "pdf_text_parser", "active", "Text extraction only; no raster measurement."),
            ParserCapability("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx_text_parser", "active", "Paragraph text extraction."),
            ParserCapability("model/ifc", "ifc_text_parser", "preview", "IFC quantities are review-gated."),
            ParserCapability("image/*", "image_ocr_queue", "blocked_without_calibration", "OCR/raster measurements require explicit calibration."),
        )

    def parse(
        self,
        *,
        document_id: str,
        filename: str,
        media_type: str,
        content: bytes,
    ) -> tuple[str, str, str, tuple[DocumentPage, ...], tuple[DocumentFact, ...], dict[str, Any]]:
        parsed = self.parse_document(
            document_id=document_id,
            filename=filename,
            media_type=media_type,
            content=content,
        )
        return parsed.media_type, parsed.parser_name, parsed.parse_status, parsed.pages, parsed.facts, parsed.metadata

    def parse_document(
        self,
        *,
        document_id: str,
        filename: str,
        media_type: str,
        content: bytes,
    ) -> ParsedDocument:
        normalized_media_type = _normalize_media_type(media_type, filename)
        parser = self.registry.select(media_type=normalized_media_type, filename=filename)
        return parser.parse(
            document_id=document_id,
            filename=filename,
            media_type=normalized_media_type,
            content=content,
        )


class InMemoryDocumentLibrary:
    """Small in-memory document library for PR8 contract development and live demos."""

    def __init__(
        self,
        parser: DocumentParser | None = None,
        *,
        embedding_config: EmbeddingConfig | None = None,
    ) -> None:
        self.parser = parser or DocumentParser()
        self.embedding_config = embedding_config
        self.documents: dict[str, DocumentRecord] = {}
        self.pages: dict[str, tuple[DocumentPage, ...]] = {}
        self.chunks: dict[str, tuple[DocumentChunk, ...]] = {}
        self.facts: dict[str, tuple[DocumentFact, ...]] = {}

    def upload(
        self,
        *,
        org_id: str,
        project_id: str,
        user_id: str,
        filename: str,
        media_type: str,
        content: bytes,
    ) -> DocumentParseResult:
        if not content:
            raise ValueError("document content is required")
        if len(content) > MAX_DOCUMENT_BYTES:
            raise ValueError("document exceeds the 15 MB parser limit")
        safe_filename = _safe_filename(filename)
        document_id = f"doc_{uuid4().hex}"
        sha256 = hashlib.sha256(content).hexdigest()
        parsed_media_type, parser_name, parse_status, pages, facts, metadata = self.parser.parse(
            document_id=document_id,
            filename=safe_filename,
            media_type=media_type,
            content=content,
        )
        record = DocumentRecord(
            id=document_id,
            org_id=org_id,
            project_id=project_id,
            filename=safe_filename,
            media_type=parsed_media_type,
            size_bytes=len(content),
            sha256=sha256,
            parser_name=parser_name,
            parser_version=self.parser.version,
            parse_status=parse_status,
            uploaded_by_user_id=user_id,
            uploaded_at=datetime.now(UTC),
            metadata=metadata,
        )
        chunks = build_document_chunks(
            document_id=document_id,
            pages=pages,
            embedding_config=self.embedding_config,
        )
        self.documents[document_id] = record
        self.pages[document_id] = pages
        self.chunks[document_id] = chunks
        self.facts[document_id] = facts
        return DocumentParseResult(record, pages, chunks, facts)

    def get_document(self, document_id: str) -> DocumentRecord:
        try:
            return self.documents[document_id]
        except KeyError as exc:
            raise DocumentNotFoundError(document_id) from exc

    def get_pages(self, document_id: str) -> tuple[DocumentPage, ...]:
        self.get_document(document_id)
        return self.pages.get(document_id, ())

    def get_facts(self, document_id: str) -> tuple[DocumentFact, ...]:
        self.get_document(document_id)
        return self.facts.get(document_id, ())

    def get_chunks(self, document_id: str) -> tuple[DocumentChunk, ...]:
        self.get_document(document_id)
        return self.chunks.get(document_id, ())

    def search_chunks(
        self,
        query: str,
        *,
        project_id: str | None = None,
        document_id: str | None = None,
        limit: int = 8,
    ) -> tuple[DocumentChunkSearchHit, ...]:
        if document_id is not None:
            candidate_chunks = self.get_chunks(document_id)
        else:
            candidate_chunks = tuple(
                chunk
                for record in self.documents.values()
                if project_id is None or record.project_id == project_id
                for chunk in self.chunks.get(record.id, ())
            )
        return search_document_chunks(
            candidate_chunks,
            query,
            limit=limit,
            embedding_config=self.embedding_config,
        )

    def review_fact(
        self,
        *,
        document_id: str,
        fact_id: str,
        review_status: DocumentReviewStatus,
        reviewed_by: str,
        note: str | None = None,
    ) -> DocumentFact:
        facts = list(self.get_facts(document_id))
        for index, fact in enumerate(facts):
            if fact.id == fact_id:
                reviewed = DocumentFact(
                    id=fact.id,
                    document_id=fact.document_id,
                    fact_type=fact.fact_type,
                    label=fact.label,
                    value=fact.value,
                    numeric_value=fact.numeric_value,
                    unit=fact.unit,
                    confidence=fact.confidence,
                    review_status=review_status,
                    source=fact.source,
                    metadata={
                        **fact.metadata,
                        "reviewed_by": reviewed_by,
                        "review_note": note,
                        "reviewed_at": datetime.now(UTC).isoformat(),
                    },
                )
                facts[index] = reviewed
                self.facts[document_id] = tuple(facts)
                return reviewed
        raise DocumentNotFoundError(fact_id)


def sample_parser_accuracy_report() -> dict[str, Any]:
    """Run the fixed canary pack through the parser and report deterministic coverage."""

    library = InMemoryDocumentLibrary()
    result = library.upload(
        org_id="org_demo_accuracy",
        project_id="project_demo_accuracy",
        user_id="system_demo_accuracy",
        filename="m1-canary-site-plan.txt",
        media_type="text/plain",
        content=SAMPLE_PACK_CONTENT,
    )
    extracted = {
        fact.label: {"numeric_value": fact.numeric_value, "unit": fact.unit}
        for fact in result.facts
    }
    matched: list[str] = []
    missing: list[str] = []
    mismatched: list[dict[str, Any]] = []
    for label, (expected_value, expected_unit) in SAMPLE_EXPECTED_FACTS.items():
        actual = extracted.get(label)
        if actual is None:
            missing.append(label)
            continue
        actual_value = actual["numeric_value"]
        actual_unit = actual["unit"]
        if (
            isinstance(actual_value, int | float)
            and abs(float(actual_value) - expected_value) < 0.0001
            and actual_unit == expected_unit
        ):
            matched.append(label)
        else:
            mismatched.append(
                {
                    "label": label,
                    "expected": {"numeric_value": expected_value, "unit": expected_unit},
                    "actual": actual,
                }
            )
    expected_count = len(SAMPLE_EXPECTED_FACTS)
    extracted_count = len(extracted)
    matched_count = len(matched)
    return {
        "demo_fixture_status": "passed" if matched_count == expected_count and not mismatched else "failed",
        "beta_status": "not_beta_ready",
        "reason": "The canary parser check passed, but real-project beta needs a broader fixture set and automated validation workflow.",
        "expected_fact_count": expected_count,
        "extracted_fact_count": extracted_count,
        "matched_fact_count": matched_count,
        "recall": matched_count / expected_count,
        "precision": matched_count / extracted_count if extracted_count else 0.0,
        "matched": matched,
        "missing": missing,
        "mismatched": mismatched,
        "blocked_for_beta": [
            "real PDF/DOCX/DXF/IFC fixture pack",
            "per-field precision/recall report across real samples",
            "automated validation workflow connected to persistence",
            "no image/PDF/raster measurement without explicit calibration",
        ],
    }


def _normalize_media_type(media_type: str, filename: str) -> str:
    lowered = (media_type or "").split(";")[0].strip().lower()
    suffix = PurePath(filename).suffix.lower()
    if lowered:
        return lowered
    return {
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".md": "text/markdown",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".dxf": "application/dxf",
        ".ifc": "model/ifc",
    }.get(suffix, "application/octet-stream")


def _parser_name(media_type: str, filename: str) -> str:
    suffix = PurePath(filename).suffix.lower()
    if media_type in DXF_MEDIA_TYPES or suffix == ".dxf":
        return "draftcheck.dxf_text_parser"
    if media_type in PDF_MEDIA_TYPES or suffix == ".pdf":
        return "draftcheck.pdf_text_parser"
    if media_type in DOCX_MEDIA_TYPES or suffix == ".docx":
        return "draftcheck.docx_text_parser"
    if media_type in IFC_MEDIA_TYPES or suffix == ".ifc":
        return "draftcheck.ifc_text_parser"
    if media_type.startswith(IMAGE_PREFIXES):
        return "draftcheck.image_ocr_queue"
    if media_type == "text/csv" or suffix == ".csv":
        return "draftcheck.csv_text_parser"
    return "draftcheck.plain_text_parser"


def _extract_text(media_type: str, filename: str, content: bytes) -> tuple[str, list[str]]:
    suffix = PurePath(filename).suffix.lower()
    notes: list[str] = []
    if media_type in PDF_MEDIA_TYPES or suffix == ".pdf":
        text = _extract_pdf_text(content)
        notes.append("pdf_text_only_no_raster_measurements")
        return text, notes
    if media_type in DOCX_MEDIA_TYPES or suffix == ".docx":
        text = _extract_docx_text(content)
        notes.append("docx_paragraph_text")
        return text, notes
    if media_type.startswith(IMAGE_PREFIXES):
        notes.append("image_ocr_requires_calibration")
        return "", notes
    decoded = _decode_text(content)
    if media_type == "text/csv" or suffix == ".csv":
        notes.append("csv_rows_flattened")
        return _flatten_csv(decoded), notes
    if media_type in IFC_MEDIA_TYPES or suffix == ".ifc":
        notes.append("ifc_quantity_preview_review_required")
    if media_type in DXF_MEDIA_TYPES or suffix == ".dxf":
        notes.append("dxf_text_entities_review_required")
    return decoded, notes


def _artifacts(
    *,
    parser_name: str,
    media_type: str,
    extraction_notes: list[str],
    pages: tuple[DocumentPage, ...],
    facts: tuple[DocumentFact, ...],
) -> tuple[DocumentArtifact, ...]:
    return (
        DocumentArtifact(
            kind="parser_output",
            label=parser_name,
            media_type=media_type,
            metadata={
                "parser_name": parser_name,
                "extraction_notes": extraction_notes,
                "page_count": len(pages),
                "fact_count": len(facts),
                "persistence_status": "in_memory_descriptor_pending_artifact_rows",
            },
        ),
    )


def extract_pdf_page_layouts(content: bytes) -> list[PdfPageExtraction]:
    """Per-page PDF text plus vector text block bboxes.

    Bounding boxes are evidence for human review only. They are not calibrated
    measurements and must not be promoted into compliance facts by this parser.
    """
    try:
        return _extract_pdf_page_layouts_with_pymupdf(content)
    except Exception as pymupdf_exc:
        try:
            texts = _extract_pdf_pages_with_pypdf(content)
        except DocumentParseError:
            raise DocumentParseError(f"Failed to parse PDF: {pymupdf_exc}") from pymupdf_exc
        return [
            PdfPageExtraction(
                page_number=index,
                text=text,
                width=None,
                height=None,
                rotation_degrees=None,
                text_blocks=(),
                vector_paths=(),
                extraction_method="pypdf_text_layer",
            )
            for index, text in enumerate(texts, start=1)
        ]


def extract_pdf_pages(content: bytes) -> list[str]:
    """Per-page text from a PDF. Raises DocumentParseError on unreadable input."""
    return [page.text for page in extract_pdf_page_layouts(content)]


def _extract_pdf_page_layouts_with_pymupdf(content: bytes) -> list[PdfPageExtraction]:
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        pages: list[PdfPageExtraction] = []
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            blocks: list[PdfTextBlock] = []
            for block in page.get_text("blocks") or []:
                if len(block) < 5:
                    continue
                block_text = str(block[4]).strip()
                if not block_text:
                    continue
                block_type = int(block[6]) if len(block) > 6 and isinstance(block[6], int) else 0
                if block_type != 0:
                    continue
                blocks.append(
                    PdfTextBlock(
                        text=block_text,
                        bbox=(
                            round(float(block[0]), 3),
                            round(float(block[1]), 3),
                            round(float(block[2]), 3),
                            round(float(block[3]), 3),
                        ),
                        block_number=int(block[5]) if len(block) > 5 else None,
                    )
                )
            rect = page.rect
            vector_paths = _pdf_vector_paths(page)
            pages.append(
                PdfPageExtraction(
                    page_number=page_index,
                    text=text,
                    width=round(float(rect.width), 3),
                    height=round(float(rect.height), 3),
                    rotation_degrees=float(page.rotation or 0),
                    text_blocks=tuple(blocks),
                    vector_paths=vector_paths,
                    extraction_method="pymupdf_text_blocks",
                )
            )
        return pages
    finally:
        doc.close()


def _pdf_vector_paths(page: Any) -> tuple[PdfVectorPath, ...]:
    get_drawings = getattr(page, "get_drawings", None)
    if not callable(get_drawings):
        return ()
    try:
        drawings = get_drawings() or []
    except Exception:
        return ()

    paths: list[PdfVectorPath] = []
    for drawing in drawings:
        rect = drawing.get("rect") if isinstance(drawing, dict) else None
        items = drawing.get("items", ()) if isinstance(drawing, dict) else ()
        drawing_type = drawing.get("type") if isinstance(drawing, dict) else None
        paths.append(
            PdfVectorPath(
                bbox=_pdf_rect_bbox(rect),
                item_count=len(items) if hasattr(items, "__len__") else 0,
                drawing_type=str(drawing_type) if drawing_type is not None else None,
            )
        )
    return tuple(paths)


def _pdf_rect_bbox(rect: Any) -> tuple[float, float, float, float] | None:
    if rect is None:
        return None
    try:
        return (
            round(float(rect.x0), 3),
            round(float(rect.y0), 3),
            round(float(rect.x1), 3),
            round(float(rect.y1), 3),
        )
    except (AttributeError, TypeError, ValueError):
        return None


def _extract_pdf_pages_with_pypdf(content: bytes) -> list[str]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        return [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse PDF: {exc}") from exc


def extract_docx_text(content: bytes) -> str:
    """Paragraph text from a DOCX. Raises DocumentParseError on unreadable input."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX: {exc}") from exc
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def decode_text_bytes(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_pdf_text(content: bytes) -> str:
    return "\n\n".join(extract_pdf_pages(content))


def _extract_docx_text(content: bytes) -> str:
    return extract_docx_text(content)


_decode_text = decode_text_bytes


def _flatten_csv(text: str) -> str:
    rows = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def _pages(
    document_id: str,
    text: str,
    parser_name: str,
    parser_version: str,
    notes: list[str],
) -> tuple[DocumentPage, ...]:
    if not text.strip():
        return ()
    chunks = [chunk.strip() for chunk in re.split(r"\f+|\n\s*---page---\s*\n", text) if chunk.strip()]
    if not chunks:
        chunks = [text.strip()]
    return tuple(
        DocumentPage(
            id=f"page_{document_id}_{index}",
            document_id=document_id,
            page_number=index,
            text=chunk,
            parser_name=parser_name,
            parser_version=parser_version,
            metadata={"notes": notes},
        )
        for index, chunk in enumerate(chunks, start=1)
    )


def _facts(document_id: str, text: str, parser_name: str, media_type: str) -> tuple[DocumentFact, ...]:
    facts: list[DocumentFact] = []
    lower_text = text.lower()
    for label, patterns, unit in (
        ("lot area", (r"lot\s+area[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m2|sqm)?",), "m2"),
        ("building footprint", (r"(?:building\s+)?footprint[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m2|sqm)?",), "m2"),
        ("open space", (r"open\s+space[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m2|sqm)?",), "m2"),
        ("front setback", (r"(?:front|primary)\s+setback[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m|metres?)?",), "m"),
        ("garage width", (r"garage\s+width[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m|metres?)?",), "m"),
        ("boundary wall length", (r"boundary\s+wall(?:\s+length)?[:\s]+([0-9]+(?:\.[0-9]+)?)\s*(m|metres?)?",), "m"),
    ):
        match = _first_match(patterns, lower_text)
        if match:
            value = float(match.group(1))
            facts.append(
                _fact(
                    document_id=document_id,
                    label=label,
                    fact_type="drawing_measurement",
                    value=value,
                    unit=unit,
                    parser_name=parser_name,
                    method="text_measurement_pattern",
                    confidence=0.72,
                )
            )
    facts.extend(_title_block_facts(document_id, text, parser_name))
    facts.extend(_dxf_facts(document_id, text, parser_name))
    facts.extend(_ifc_facts(document_id, text, parser_name, media_type))
    return tuple(_dedupe_facts(facts))


def _first_match(patterns: tuple[str, ...], text: str) -> re.Match[str] | None:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match
    return None


TITLE_BLOCK_PATTERNS: tuple[tuple[str, str, tuple[str, ...]], ...] = (
    (
        "drawing title",
        "drawing_title",
        (
            r"^\s*(?:drawing\s+)?title\s*[:|]\s*(?P<value>.+?)\s*$",
            r"^\s*sheet\s+title\s*[:|]\s*(?P<value>.+?)\s*$",
        ),
    ),
    (
        "drawing number",
        "drawing_number",
        (
            r"^\s*(?:drawing|dwg|sheet)\s*(?:no\.?|number)\s*[:|]\s*(?P<value>[A-Za-z0-9._/-]+)\s*$",
            r"^\s*(?:drawing|dwg)\s*#\s*[:|]?\s*(?P<value>[A-Za-z0-9._/-]+)\s*$",
        ),
    ),
    (
        "revision",
        "revision",
        (
            r"^\s*(?:rev(?:ision)?)\s*[:|]\s*(?P<value>[A-Za-z0-9._/-]+)\s*$",
        ),
    ),
    (
        "scale",
        "scale",
        (
            r"^\s*scale\s*[:|]\s*(?P<value>(?:1\s*:\s*\d+|[A-Za-z0-9 .:/-]+))\s*$",
        ),
    ),
    (
        "project number",
        "project_number",
        (
            r"^\s*(?:project|job)\s*(?:no\.?|number)\s*[:|]\s*(?P<value>[A-Za-z0-9._/-]+)\s*$",
        ),
    ),
)


def _title_block_facts(document_id: str, text: str, parser_name: str) -> list[DocumentFact]:
    facts: list[DocumentFact] = []
    seen_fields: set[str] = set()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or len(line) > 240:
            continue
        for label, field, patterns in TITLE_BLOCK_PATTERNS:
            if field in seen_fields:
                continue
            match = _first_match(patterns, line)
            if match is None:
                continue
            value = _clean_title_block_value(match.group("value"))
            if not value:
                continue
            facts.append(
                _metadata_fact(
                    document_id=document_id,
                    label=label,
                    fact_type="title_block_metadata",
                    value=value,
                    parser_name=parser_name,
                    method="title_block_text_pattern",
                    confidence=0.66,
                    metadata={
                        "title_block_field": field,
                        "source_text": raw_line.strip(),
                        "metadata_compliance_ready": False,
                        "metadata_readiness_reason": "human confirmation required before use",
                    },
                )
            )
            seen_fields.add(field)
    return facts


def _clean_title_block_value(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value.replace("\x00", " ")).strip(" -:|")
    return cleaned[:160]


def _dxf_facts(document_id: str, text: str, parser_name: str) -> list[DocumentFact]:
    unit_context = _dxf_unit_context(text)
    candidates = _extract_dxf_dimension_candidates(text)
    if candidates:
        facts: list[DocumentFact] = []
        for index, candidate in enumerate(candidates, start=1):
            value, unit, unit_metadata = _normalise_dxf_measurement(
                candidate.measurement,
                unit_context,
            )
            base_confidence = 0.62 if candidate.insert_scale_uncertain else 0.76
            facts.append(
                _fact(
                    document_id=document_id,
                    label=f"dxf dimension {index}",
                    fact_type="drawing_dimension",
                    value=value,
                    unit=unit,
                    parser_name=parser_name,
                    method=candidate.extraction_method,
                    confidence=_cap_confidence(base_confidence, unit_context),
                    metadata={**_dxf_dimension_metadata(candidate), **unit_metadata},
                )
            )
        return facts

    facts: list[DocumentFact] = []
    if "a-dimensions" not in text.lower() and "dimension" not in text.lower():
        return facts
    for index, value in enumerate(
        re.findall(
            r"\b(?:DIMENSION|LINE_LENGTH|MEASUREMENT)[:=\s]+([0-9]+(?:\.[0-9]+)?)",
            text,
            flags=re.I,
        ),
        start=1,
    ):
        value, unit, unit_metadata = _normalise_dxf_measurement(float(value), unit_context)
        facts.append(
            _fact(
                document_id=document_id,
                label=f"dxf dimension {index}",
                fact_type="drawing_dimension",
                value=value,
                unit=unit,
                parser_name=parser_name,
                method="dxf_text_dimension_entity",
                confidence=_cap_confidence(0.68, unit_context),
                metadata={
                    "entity_type": "DIMENSION",
                    "cad_extraction_review_status": "pending_review",
                    "scale_status": "not_available_from_text_fallback",
                    **unit_metadata,
                },
            )
        )
    return facts


def _extract_dxf_dimension_candidates(text: str) -> list[_DxfDimensionCandidate]:
    candidates = _extract_dxf_dimensions_with_ezdxf(text)
    if candidates:
        return candidates
    return _extract_dxf_dimensions_from_group_codes(text)


def _extract_dxf_dimensions_with_ezdxf(text: str) -> list[_DxfDimensionCandidate]:
    try:
        import ezdxf  # type: ignore[import-not-found]
    except ImportError:
        return []

    try:
        doc = ezdxf.read(io.StringIO(text))
    except Exception:
        return []

    candidates: list[_DxfDimensionCandidate] = []
    for entity in doc.modelspace():
        entity_type = str(entity.dxftype()).upper()
        if entity_type == "DIMENSION":
            candidate = _ezdxf_dimension_candidate(entity)
            if candidate is not None:
                candidates.append(candidate)
        elif entity_type == "INSERT":
            candidates.extend(_ezdxf_insert_dimension_candidates(doc, entity))
    return candidates


def _ezdxf_dimension_candidate(entity: Any) -> _DxfDimensionCandidate | None:
    measurement = _ezdxf_dimension_measurement(entity)
    if measurement is None:
        return None
    text_override = _clean_dxf_text_override(getattr(entity.dxf, "text", None))
    override_value = _first_number(text_override or "")
    return _DxfDimensionCandidate(
        measurement=measurement,
        handle=str(getattr(entity.dxf, "handle", "") or "") or None,
        layer=str(getattr(entity.dxf, "layer", "") or "") or None,
        entity_type=str(entity.dxftype()).upper(),
        cad_space=_dxf_cad_space_from_flag(getattr(entity.dxf, "paperspace", None)),
        text_override=text_override,
        text_override_numeric_value=override_value,
        text_override_differs=_dxf_override_differs(measurement, override_value),
        extraction_method="ezdxf_dimension_entity",
    )


def _ezdxf_insert_dimension_candidates(doc: Any, insert: Any) -> list[_DxfDimensionCandidate]:
    block_name = str(getattr(insert.dxf, "name", "") or "")
    if not block_name:
        return []
    try:
        block = doc.blocks.get(block_name)
    except Exception:
        return []

    sx = _float_or_none(getattr(insert.dxf, "xscale", 1.0)) or 1.0
    sy = _float_or_none(getattr(insert.dxf, "yscale", 1.0)) or 1.0
    sz = _float_or_none(getattr(insert.dxf, "zscale", 1.0)) or 1.0
    scale_uncertain = not _nearly_equal(abs(sx), abs(sy))
    scale_factor = abs(sx) if not scale_uncertain else 1.0
    insert_handle = str(getattr(insert.dxf, "handle", "") or "") or None
    insert_layer = str(getattr(insert.dxf, "layer", "") or "") or None
    insert_cad_space = _dxf_cad_space_from_flag(getattr(insert.dxf, "paperspace", None))

    candidates: list[_DxfDimensionCandidate] = []
    for entity in block:
        if str(entity.dxftype()).upper() != "DIMENSION":
            continue
        base = _ezdxf_dimension_candidate(entity)
        if base is None:
            continue
        candidates.append(
            _DxfDimensionCandidate(
                measurement=base.measurement * scale_factor,
                handle=base.handle,
                layer=base.layer,
                entity_type=base.entity_type,
                text_override=base.text_override,
                text_override_numeric_value=base.text_override_numeric_value,
                text_override_differs=base.text_override_differs,
                block_name=block_name,
                cad_space=insert_cad_space or base.cad_space,
                layout_name=base.layout_name,
                insert_handle=insert_handle,
                insert_layer=insert_layer,
                insert_scale_x=sx,
                insert_scale_y=sy,
                insert_scale_z=sz,
                insert_scale_applied=not scale_uncertain and not _nearly_equal(scale_factor, 1.0),
                insert_scale_uncertain=scale_uncertain,
                scale_review_reason=(
                    "non_uniform_insert_scale_requires_operator_review" if scale_uncertain else None
                ),
                extraction_method="ezdxf_block_insert_dimension_entity",
            )
        )
    return candidates


def _ezdxf_dimension_measurement(entity: Any) -> float | None:
    for attr_name in ("get_measurement", "actual_measurement"):
        attr = getattr(entity, attr_name, None)
        try:
            value = attr() if callable(attr) else attr
        except Exception:
            continue
        numeric_value = _float_or_none(value)
        if numeric_value is not None:
            return numeric_value
    return _float_or_none(getattr(entity.dxf, "actual_measurement", None))


def _extract_dxf_dimensions_from_group_codes(text: str) -> list[_DxfDimensionCandidate]:
    pairs = _dxf_group_pairs(text)
    if not pairs:
        return []

    modelspace_dimensions: list[_DxfDimensionCandidate] = []
    block_dimensions: dict[str, list[_DxfDimensionCandidate]] = {}
    inserts: list[list[tuple[str, str]]] = []
    section: str | None = None
    current_block_name: str | None = None
    index = 0
    while index < len(pairs):
        code, value = pairs[index]
        if code != "0":
            index += 1
            continue

        entity_type = value.strip().upper()
        entity_pairs, next_index = _dxf_entity_pairs(pairs, index)
        if entity_type == "SECTION":
            section = _dxf_first(entity_pairs, "2")
        elif entity_type == "ENDSEC":
            section = None
        elif section == "BLOCKS" and entity_type == "BLOCK":
            current_block_name = _dxf_first(entity_pairs, "2")
            if current_block_name:
                block_dimensions.setdefault(current_block_name, [])
        elif section == "BLOCKS" and entity_type == "ENDBLK":
            current_block_name = None
        elif entity_type == "DIMENSION":
            candidate = _dxf_dimension_from_pairs(entity_pairs)
            if candidate is not None:
                if section == "BLOCKS" and current_block_name:
                    block_dimensions.setdefault(current_block_name, []).append(
                        _replace_dxf_candidate(candidate, block_name=current_block_name)
                    )
                elif section == "ENTITIES":
                    modelspace_dimensions.append(candidate)
        elif section == "ENTITIES" and entity_type == "INSERT":
            inserts.append(entity_pairs)
        index = next_index

    candidates = list(modelspace_dimensions)
    for insert_pairs in inserts:
        candidates.extend(_dxf_block_insert_candidates(insert_pairs, block_dimensions))
    return candidates


def _dxf_group_pairs(text: str) -> list[tuple[str, str]]:
    lines = [line.strip() for line in text.splitlines()]
    pairs: list[tuple[str, str]] = []
    index = 0
    while index + 1 < len(lines):
        code = lines[index]
        value = lines[index + 1]
        if re.fullmatch(r"-?\d+", code):
            pairs.append((code, value))
            index += 2
        else:
            index += 1
    return pairs


DXF_INSUNITS: dict[int, tuple[str, float | None]] = {
    0: ("unitless", None),
    1: ("inches", 0.0254),
    2: ("feet", 0.3048),
    4: ("millimetres", 0.001),
    5: ("centimetres", 0.01),
    6: ("metres", 1.0),
    7: ("kilometres", 1000.0),
}


def _dxf_unit_context(text: str) -> _DxfUnitContext:
    pairs = _dxf_group_pairs(text)
    for index, (code, value) in enumerate(pairs):
        if code != "9" or value.strip().upper() != "$INSUNITS":
            continue
        for next_code, next_value in pairs[index + 1 : index + 4]:
            if next_code != "70":
                continue
            unit_code = _int_or_none(next_value)
            if unit_code is None:
                break
            declared_units, scale = DXF_INSUNITS.get(unit_code, (f"insunits_{unit_code}", None))
            return _unit_context_from_code(unit_code, declared_units, scale)
    return _unit_context_from_code(None, "missing", None)


def _unit_context_from_code(
    unit_code: int | None,
    declared_units: str,
    scale_to_metres: float | None,
) -> _DxfUnitContext:
    if scale_to_metres is None:
        return _DxfUnitContext(
            code=unit_code,
            declared_units=declared_units,
            output_unit="drawing_unit",
            scale_to_metres=None,
            confidence_cap=0.4,
        )
    return _DxfUnitContext(
        code=unit_code,
        declared_units=declared_units,
        output_unit="m",
        scale_to_metres=scale_to_metres,
        confidence_cap=None,
    )


def _normalise_dxf_measurement(
    value: float,
    unit_context: _DxfUnitContext,
) -> tuple[float, str, dict[str, Any]]:
    metadata: dict[str, Any] = {
        "dxf_insunits_code": unit_context.code,
        "dxf_declared_units": unit_context.declared_units,
        "unit_declared": unit_context.unit_declared,
        "unit_conversion_applied": False,
    }
    if unit_context.scale_to_metres is None:
        metadata["measurement_readiness_reason"] = (
            "DXF $INSUNITS is missing, unitless, or unsupported; operator must confirm drawing units"
        )
        return value, unit_context.output_unit, metadata

    normalised = value * unit_context.scale_to_metres
    metadata["unit_conversion_applied"] = not _nearly_equal(unit_context.scale_to_metres, 1.0)
    if metadata["unit_conversion_applied"]:
        metadata["source_numeric_value"] = value
        metadata["source_unit"] = unit_context.declared_units
    return normalised, unit_context.output_unit, metadata


def _cap_confidence(base_confidence: float, unit_context: _DxfUnitContext) -> float:
    if unit_context.confidence_cap is None:
        return base_confidence
    return min(base_confidence, unit_context.confidence_cap)


def _dxf_entity_pairs(
    pairs: list[tuple[str, str]],
    start_index: int,
) -> tuple[list[tuple[str, str]], int]:
    end_index = start_index + 1
    while end_index < len(pairs) and pairs[end_index][0] != "0":
        end_index += 1
    return pairs[start_index:end_index], end_index


def _dxf_dimension_from_pairs(entity_pairs: list[tuple[str, str]]) -> _DxfDimensionCandidate | None:
    measurement = _float_or_none(_dxf_first(entity_pairs, "42"))
    if measurement is None:
        return None
    text_override = _clean_dxf_text_override(_dxf_first(entity_pairs, "1"))
    override_value = _first_number(text_override or "")
    return _DxfDimensionCandidate(
        measurement=measurement,
        handle=_dxf_first(entity_pairs, "5"),
        layer=_dxf_first(entity_pairs, "8"),
        entity_type="DIMENSION",
        cad_space=_dxf_cad_space_from_flag(_dxf_first(entity_pairs, "67")),
        layout_name=_dxf_first(entity_pairs, "410"),
        text_override=text_override,
        text_override_numeric_value=override_value,
        text_override_differs=_dxf_override_differs(measurement, override_value),
    )


def _dxf_block_insert_candidates(
    insert_pairs: list[tuple[str, str]],
    block_dimensions: dict[str, list[_DxfDimensionCandidate]],
) -> list[_DxfDimensionCandidate]:
    block_name = _dxf_first(insert_pairs, "2")
    if not block_name:
        return []
    base_dimensions = block_dimensions.get(block_name, [])
    if not base_dimensions:
        return []

    sx = _float_or_none(_dxf_first(insert_pairs, "41")) or 1.0
    sy = _float_or_none(_dxf_first(insert_pairs, "42")) or 1.0
    sz = _float_or_none(_dxf_first(insert_pairs, "43")) or 1.0
    scale_uncertain = not _nearly_equal(abs(sx), abs(sy))
    scale_factor = abs(sx) if not scale_uncertain else 1.0
    return [
        _DxfDimensionCandidate(
            measurement=dimension.measurement * scale_factor,
            handle=dimension.handle,
            layer=dimension.layer,
            entity_type=dimension.entity_type,
            text_override=dimension.text_override,
            text_override_numeric_value=dimension.text_override_numeric_value,
            text_override_differs=dimension.text_override_differs,
            block_name=block_name,
            cad_space=_dxf_cad_space_from_flag(_dxf_first(insert_pairs, "67")) or dimension.cad_space,
            layout_name=_dxf_first(insert_pairs, "410") or dimension.layout_name,
            insert_handle=_dxf_first(insert_pairs, "5"),
            insert_layer=_dxf_first(insert_pairs, "8"),
            insert_scale_x=sx,
            insert_scale_y=sy,
            insert_scale_z=sz,
            insert_scale_applied=not scale_uncertain and not _nearly_equal(scale_factor, 1.0),
            insert_scale_uncertain=scale_uncertain,
            scale_review_reason=(
                "non_uniform_insert_scale_requires_operator_review" if scale_uncertain else None
            ),
            extraction_method="dxf_group_code_block_insert_dimension_entity",
        )
        for dimension in base_dimensions
    ]


def _dxf_first(entity_pairs: list[tuple[str, str]], code: str) -> str | None:
    for pair_code, value in entity_pairs:
        if pair_code == code:
            return value.strip() or None
    return None


def _dxf_cad_space_from_flag(value: object) -> str | None:
    if value is None:
        return None
    flag = str(value).strip()
    if flag == "0":
        return "model_space"
    if flag == "1":
        return "paper_space"
    return None


def _replace_dxf_candidate(
    candidate: _DxfDimensionCandidate,
    *,
    block_name: str | None,
) -> _DxfDimensionCandidate:
    return _DxfDimensionCandidate(
        measurement=candidate.measurement,
        handle=candidate.handle,
        layer=candidate.layer,
        entity_type=candidate.entity_type,
        text_override=candidate.text_override,
        text_override_numeric_value=candidate.text_override_numeric_value,
        text_override_differs=candidate.text_override_differs,
        block_name=block_name,
        cad_space=candidate.cad_space,
        layout_name=candidate.layout_name,
        extraction_method=candidate.extraction_method,
    )


def _dxf_dimension_metadata(candidate: _DxfDimensionCandidate) -> dict[str, Any]:
    metadata: dict[str, Any] = {
        "entity_type": candidate.entity_type,
        "entity_handle": candidate.handle,
        "entity_layer": candidate.layer,
        "cad_space": candidate.cad_space,
        "layout_name": candidate.layout_name,
        "cad_extraction_review_status": "pending_review",
        "scale_status": "not_insert_scaled",
    }
    if candidate.text_override:
        metadata.update(
            {
                "text_override": candidate.text_override,
                "text_override_numeric_value": candidate.text_override_numeric_value,
                "text_override_differs": candidate.text_override_differs,
            }
        )
    if candidate.block_name:
        scale_status = "insert_scale_uncertain" if candidate.insert_scale_uncertain else "insert_scale_known"
        metadata.update(
            {
                "block_name": candidate.block_name,
                "insert_handle": candidate.insert_handle,
                "insert_layer": candidate.insert_layer,
                "insert_scale": {
                    "x": candidate.insert_scale_x,
                    "y": candidate.insert_scale_y,
                    "z": candidate.insert_scale_z,
                },
                "insert_scale_applied": candidate.insert_scale_applied,
                "insert_scale_uncertain": candidate.insert_scale_uncertain,
                "scale_status": scale_status,
                "scale_review_reason": candidate.scale_review_reason,
            }
        )
    return metadata


def _clean_dxf_text_override(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text or text == "<>":
        return None
    return text


def _dxf_override_differs(measurement: float, override_value: float | None) -> bool:
    return override_value is not None and not _nearly_equal(measurement, override_value)


def _first_number(text: str) -> float | None:
    match = re.search(r"-?[0-9]+(?:\.[0-9]+)?", text)
    return _float_or_none(match.group(0)) if match else None


def _float_or_none(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _int_or_none(value: object) -> int | None:
    if value is None:
        return None
    try:
        return int(str(value).strip())
    except (TypeError, ValueError):
        return None


def _nearly_equal(left: float, right: float, *, tolerance: float = 0.000001) -> bool:
    return abs(left - right) <= tolerance


def _ifc_facts(document_id: str, text: str, parser_name: str, media_type: str) -> list[DocumentFact]:
    if media_type not in IFC_MEDIA_TYPES and ".ifc" not in text[:200].lower():
        return []
    facts: list[DocumentFact] = []
    for index, quantity in enumerate(_ifc_quantity_candidates(text), start=1):
        related_object = quantity.get("related_object")
        parser_status = str(quantity.get("parser_status") or "step_text_fallback")
        facts.append(
            _fact(
                document_id=document_id,
                label=f"ifc {quantity['quantity_type']} quantity {index}",
                fact_type="model_quantity",
                value=quantity["value"],
                unit=quantity["unit"],
                parser_name=parser_name,
                method=str(quantity.get("method") or "ifc_step_quantity_preview"),
                confidence=0.78 if parser_status == "ifcopenshell" else 0.62,
                metadata={
                    "ifc_entity_id": quantity["entity_id"],
                    "ifc_quantity_type": quantity["quantity_type"],
                    "ifc_quantity_name": quantity["name"],
                    "ifc_unit_ref": quantity["unit_ref"],
                    "ifc_quantity_set_id": quantity.get("quantity_set_id"),
                    "ifc_quantity_set_name": quantity.get("quantity_set_name"),
                    "ifc_related_object_id": related_object.get("entity_id") if related_object else None,
                    "ifc_related_object_type": related_object.get("entity_type") if related_object else None,
                    "ifc_related_object_name": related_object.get("name") if related_object else None,
                    "ifc_related_object_long_name": related_object.get("long_name") if related_object else None,
                    "ifc_parser_status": parser_status,
                    "measurement_readiness_reason": "IFC quantity requires review before compliance use",
                },
            )
        )
    return facts


def _ifc_quantity_candidates(text: str) -> list[dict[str, Any]]:
    ifcopenshell_candidates = _ifcopenshell_quantity_candidates(text)
    if ifcopenshell_candidates:
        return ifcopenshell_candidates

    candidates: list[dict[str, Any]] = []
    context = _ifc_context(text)
    pattern = re.compile(
        r"#(?P<entity_id>\d+)\s*=\s*IFCQUANTITY(?P<quantity_type>AREA|LENGTH|VOLUME|COUNT)\((?P<body>.*?)\);",
        flags=re.IGNORECASE | re.DOTALL,
    )
    for match in pattern.finditer(text):
        args = _ifc_split_args(match.group("body"))
        if len(args) < 4:
            continue
        value = _float_or_none(args[3])
        if value is None:
            continue
        quantity_type = match.group("quantity_type").lower()
        entity_id = f"#{match.group('entity_id')}"
        quantity_set = context["quantity_sets_by_quantity"].get(entity_id)
        candidates.append(
            {
                "entity_id": entity_id,
                "quantity_type": quantity_type,
                "name": _ifc_unquote(args[0]),
                "unit": _ifc_default_unit(quantity_type),
                "unit_ref": None if args[2].strip() == "$" else args[2].strip(),
                "value": value,
                "quantity_set_id": quantity_set["entity_id"] if quantity_set else None,
                "quantity_set_name": quantity_set["name"] if quantity_set else None,
                "related_object": context["objects_by_quantity_set"].get(quantity_set["entity_id"])
                if quantity_set
                else None,
                "parser_status": "step_text_fallback",
                "method": "ifc_step_quantity_preview",
            }
        )
    return candidates


def _ifcopenshell_quantity_candidates(text: str) -> list[dict[str, Any]]:
    try:
        import ifcopenshell  # type: ignore[import-not-found]
    except ImportError:
        return []

    try:
        model = ifcopenshell.file.from_string(text)
    except Exception:
        return []

    quantities_by_id: dict[str, dict[str, Any]] = {}
    for entity_name, quantity_type, value_attr, unit in (
        ("IfcQuantityArea", "area", "AreaValue", "m2"),
        ("IfcQuantityLength", "length", "LengthValue", "m"),
        ("IfcQuantityVolume", "volume", "VolumeValue", "m3"),
        ("IfcQuantityCount", "count", "CountValue", "count"),
    ):
        for quantity in _ifc_by_type(model, entity_name):
            value = _float_or_none(getattr(quantity, value_attr, None))
            if value is None:
                continue
            entity_id = _ifc_runtime_entity_id(quantity)
            quantities_by_id[entity_id] = {
                "entity_id": entity_id,
                "quantity_type": quantity_type,
                "name": str(getattr(quantity, "Name", "") or ""),
                "unit": unit,
                "unit_ref": _ifc_runtime_entity_id(getattr(quantity, "Unit", None))
                if getattr(quantity, "Unit", None) is not None
                else None,
                "value": value,
                "quantity_set_id": None,
                "quantity_set_name": None,
                "related_object": None,
                "parser_status": "ifcopenshell",
                "method": "ifcopenshell_quantity_parser",
            }

    if not quantities_by_id:
        return []

    quantity_sets = _ifcopenshell_quantity_sets(model)
    for quantity_set_id, quantity_set in quantity_sets.items():
        for quantity_id in quantity_set["quantity_ids"]:
            quantity = quantities_by_id.get(quantity_id)
            if quantity is None:
                continue
            quantity["quantity_set_id"] = quantity_set_id
            quantity["quantity_set_name"] = quantity_set["name"]

    objects_by_quantity_set = _ifcopenshell_objects_by_quantity_set(model)
    for quantity in quantities_by_id.values():
        quantity_set_id = quantity.get("quantity_set_id")
        if quantity_set_id:
            quantity["related_object"] = objects_by_quantity_set.get(str(quantity_set_id))

    return list(quantities_by_id.values())


def _ifc_by_type(model: Any, entity_name: str) -> tuple[Any, ...]:
    by_type = getattr(model, "by_type", None)
    if not callable(by_type):
        return ()
    try:
        return tuple(by_type(entity_name) or ())
    except Exception:
        return ()


def _ifcopenshell_quantity_sets(model: Any) -> dict[str, dict[str, Any]]:
    quantity_sets: dict[str, dict[str, Any]] = {}
    for quantity_set in _ifc_by_type(model, "IfcElementQuantity"):
        quantity_set_id = _ifc_runtime_entity_id(quantity_set)
        quantity_sets[quantity_set_id] = {
            "entity_id": quantity_set_id,
            "name": str(getattr(quantity_set, "Name", "") or ""),
            "quantity_ids": [
                _ifc_runtime_entity_id(quantity)
                for quantity in tuple(getattr(quantity_set, "Quantities", ()) or ())
            ],
        }
    return quantity_sets


def _ifcopenshell_objects_by_quantity_set(model: Any) -> dict[str, dict[str, Any]]:
    objects_by_quantity_set: dict[str, dict[str, Any]] = {}
    for relation in _ifc_by_type(model, "IfcRelDefinesByProperties"):
        quantity_set = getattr(relation, "RelatingPropertyDefinition", None)
        quantity_set_id = _ifc_runtime_entity_id(quantity_set) if quantity_set is not None else None
        if not quantity_set_id:
            continue
        related_objects = tuple(getattr(relation, "RelatedObjects", ()) or ())
        related_object = related_objects[0] if related_objects else None
        if related_object is not None:
            objects_by_quantity_set[quantity_set_id] = _ifc_runtime_object_payload(related_object)
    return objects_by_quantity_set


def _ifc_runtime_object_payload(entity: Any) -> dict[str, Any]:
    return {
        "entity_id": _ifc_runtime_entity_id(entity),
        "entity_type": _ifc_runtime_entity_type(entity),
        "name": str(getattr(entity, "Name", "") or "") or None,
        "long_name": str(getattr(entity, "LongName", "") or "") or None,
    }


def _ifc_runtime_entity_id(entity: Any) -> str:
    if entity is None:
        return ""
    id_attr = getattr(entity, "id", None)
    try:
        value = id_attr() if callable(id_attr) else id_attr
    except Exception:
        value = None
    return f"#{value}" if value is not None else f"#{id(entity)}"


def _ifc_runtime_entity_type(entity: Any) -> str | None:
    is_a = getattr(entity, "is_a", None)
    try:
        value = is_a() if callable(is_a) else is_a
    except Exception:
        value = None
    return str(value).upper() if value else None


def _ifc_context(text: str) -> dict[str, dict[str, Any]]:
    objects = _ifc_objects(text)
    quantity_sets = _ifc_quantity_sets(text)
    quantity_sets_by_quantity: dict[str, dict[str, Any]] = {}
    for quantity_set in quantity_sets.values():
        for quantity_id in quantity_set["quantity_ids"]:
            quantity_sets_by_quantity[quantity_id] = quantity_set
    return {
        "quantity_sets_by_quantity": quantity_sets_by_quantity,
        "objects_by_quantity_set": _ifc_objects_by_quantity_set(text, objects),
    }


def _ifc_objects(text: str) -> dict[str, dict[str, Any]]:
    objects: dict[str, dict[str, Any]] = {}
    pattern = re.compile(
        r"(?P<entity_id>#\d+)\s*=\s*IFC(?P<entity_type>SPACE|SITE|BUILDING|BUILDINGSTOREY|SLAB|WALL|WALLSTANDARDCASE)\((?P<body>.*?)\);",
        flags=re.IGNORECASE | re.DOTALL,
    )
    for match in pattern.finditer(text):
        args = _ifc_split_args(match.group("body"))
        name = _ifc_optional_arg(args, 2)
        long_name = _ifc_optional_arg(args, 7)
        objects[match.group("entity_id")] = {
            "entity_id": match.group("entity_id"),
            "entity_type": f"IFC{match.group('entity_type').upper()}",
            "name": _ifc_unquote(name) if name else None,
            "long_name": _ifc_unquote(long_name) if long_name else None,
        }
    return objects


def _ifc_quantity_sets(text: str) -> dict[str, dict[str, Any]]:
    quantity_sets: dict[str, dict[str, Any]] = {}
    pattern = re.compile(
        r"(?P<entity_id>#\d+)\s*=\s*IFCELEMENTQUANTITY\((?P<body>.*?)\);",
        flags=re.IGNORECASE | re.DOTALL,
    )
    for match in pattern.finditer(text):
        args = _ifc_split_args(match.group("body"))
        quantity_sets[match.group("entity_id")] = {
            "entity_id": match.group("entity_id"),
            "name": _ifc_unquote(_ifc_optional_arg(args, 2) or ""),
            "quantity_ids": _ifc_reference_list(_ifc_optional_arg(args, 5) or ""),
        }
    return quantity_sets


def _ifc_objects_by_quantity_set(
    text: str,
    objects: dict[str, dict[str, Any]],
) -> dict[str, dict[str, Any]]:
    objects_by_quantity_set: dict[str, dict[str, Any]] = {}
    pattern = re.compile(
        r"#\d+\s*=\s*IFCRELDEFINESBYPROPERTIES\((?P<body>.*?)\);",
        flags=re.IGNORECASE | re.DOTALL,
    )
    for match in pattern.finditer(text):
        args = _ifc_split_args(match.group("body"))
        related_object_ids = _ifc_reference_list(_ifc_optional_arg(args, 4) or "")
        quantity_set_id = _ifc_optional_arg(args, 5)
        if not quantity_set_id or quantity_set_id == "$":
            continue
        related_object = next((objects[object_id] for object_id in related_object_ids if object_id in objects), None)
        if related_object is not None:
            objects_by_quantity_set[quantity_set_id] = related_object
    return objects_by_quantity_set


def _ifc_split_args(body: str) -> list[str]:
    args: list[str] = []
    current: list[str] = []
    quote_open = False
    depth = 0
    index = 0
    while index < len(body):
        char = body[index]
        if char == "'":
            quote_open = not quote_open
            current.append(char)
        elif not quote_open and char == "(":
            depth += 1
            current.append(char)
        elif not quote_open and char == ")":
            depth = max(0, depth - 1)
            current.append(char)
        elif not quote_open and depth == 0 and char == ",":
            args.append("".join(current).strip())
            current = []
        else:
            current.append(char)
        index += 1
    args.append("".join(current).strip())
    return args


def _ifc_optional_arg(args: list[str], index: int) -> str | None:
    if index >= len(args):
        return None
    value = args[index].strip()
    return None if value == "$" else value


def _ifc_reference_list(value: str) -> list[str]:
    return [match.group(0) for match in re.finditer(r"#\d+", value)]


def _ifc_unquote(value: str) -> str:
    stripped = value.strip()
    if len(stripped) >= 2 and stripped[0] == "'" and stripped[-1] == "'":
        return stripped[1:-1].replace("''", "'")
    return stripped


def _ifc_default_unit(quantity_type: str) -> str:
    return {
        "area": "m2",
        "length": "m",
        "volume": "m3",
        "count": "count",
    }.get(quantity_type, "unknown")


def _fact(
    *,
    document_id: str,
    label: str,
    fact_type: str,
    value: float,
    unit: str,
    parser_name: str,
    method: str,
    confidence: float,
    metadata: dict[str, Any] | None = None,
) -> DocumentFact:
    fact_metadata = {
        "method": method,
        "measurement_compliance_ready": False,
        "measurement_readiness_reason": "human promotion required before compliance use",
    }
    if metadata:
        fact_metadata.update(metadata)
    return DocumentFact(
        id=f"fact_{uuid4().hex}",
        document_id=document_id,
        fact_type=fact_type,
        label=label,
        value={"value": value, "unit": unit},
        numeric_value=value,
        unit=unit,
        confidence=confidence,
        review_status=DocumentReviewStatus.PENDING_REVIEW,
        source=parser_name,
        metadata=fact_metadata,
    )


def _metadata_fact(
    *,
    document_id: str,
    label: str,
    fact_type: str,
    value: str,
    parser_name: str,
    method: str,
    confidence: float,
    metadata: dict[str, Any] | None = None,
) -> DocumentFact:
    fact_metadata = {
        "method": method,
        "measurement_compliance_ready": False,
        "measurement_readiness_reason": "not a measurement",
    }
    if metadata:
        fact_metadata.update(metadata)
    return DocumentFact(
        id=f"fact_{uuid4().hex}",
        document_id=document_id,
        fact_type=fact_type,
        label=label,
        value={"value": value},
        numeric_value=None,
        unit=None,
        confidence=confidence,
        review_status=DocumentReviewStatus.PENDING_REVIEW,
        source=parser_name,
        metadata=fact_metadata,
    )


def _dedupe_facts(facts: list[DocumentFact]) -> list[DocumentFact]:
    seen: set[tuple[str, float | None, str | None]] = set()
    deduped: list[DocumentFact] = []
    for fact in facts:
        key = (fact.label, fact.numeric_value, fact.unit)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(fact)
    return deduped


def _safe_filename(filename: str) -> str:
    name = PurePath(filename or "upload.bin").name.replace("\x00", "")
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" .")
    return name or "upload.bin"
